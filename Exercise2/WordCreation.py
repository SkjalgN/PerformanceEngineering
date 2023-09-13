
import docx
import matplotlib.pyplot as plt

class WordCreation :
    
    #create a word document with the data
    def create_word_document(self,GC):
        doc = docx.Document()
        doc.add_heading('Assignment 2: Chess Games with Stockfish', 0)

        #1.1 Introduction
        doc.add_heading('1.1 Introduction', level=1)
        doc.add_paragraph('This document is made to show the results of the chess games played with Stockfish. It includes statistics, graphs and descriptions of the results.')
        
        #2.1 General games
        doc.add_heading('2.1 General games', level=1)
        doc.add_paragraph('Here you can see some general info about the games played')

        doc.add_heading('2.1.1 Statistics', level=2)
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        for row in table.rows:
            row.height =docx.shared.Cm(1)
        table.cell(1,0).text = 'Amount of games'
        table.cell(0,1).text = 'White wins'
        table.cell(0,2).text = 'Black wins'
        table.cell(0,3).text = 'Games Drawn'
        table.cell(0,4).text = 'Total games'
        table.cell(1,1).text = str(GC.white_wins())
        table.cell(1,2).text = str(GC.black_wins())
        table.cell(1,3).text = str(GC.draws())
        table.cell(1,4).text = str(GC.total_games())


        doc.add_heading('2.1.2 Graphs', level=2)
        doc.add_paragraph('Here you can see a graph showing how many games were still going after x moves.')
        plt.figure(figsize=(6, 4))
        plt.plot(GC.stockfish_move_list(), color='blue')
        plt.xlabel('Move Number')
        plt.ylabel('Remaining Games')
        plt.legend(['Total'])
        plt.savefig('Plots/my_plot.png')
        doc.add_picture('Plots/my_plot.png')
        doc.add_paragraph(' ')

        #2.2 Games focused on Stockfish
        doc.add_heading('2.2 Games focused on Stockfish', level=1)
        doc.add_paragraph('Here you can see some general info about the games Stockfish played')

        doc.add_heading('2.2.1 Statistics', level=2)
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        for row in table.rows:
            row.height =docx.shared.Cm(1)
        table.cell(1,0).text = 'Total'
        table.cell(2,0).text = 'As White'
        table.cell(3,0).text = 'As Black'
        table.cell(0,1).text = 'Games Won'
        table.cell(0,2).text = 'Games Lost'
        table.cell(0,3).text = 'Games Drawn'
        table.cell(1,1).text = str(GC.stockfish_wins())
        table.cell(1,2).text = str(GC.stockfish_losses())
        table.cell(1,3).text = str(GC.stockfish_draws())
        table.cell(2,1).text = str(GC.stockfish_wins_with_white())
        table.cell(2,2).text = str(GC.stockfish_losses_with_white())
        table.cell(2,3).text = str(GC.stockfish_draws_with_white())
        table.cell(3,1).text = str(GC.stockfish_wins_with_black())
        table.cell(3,2).text = str(GC.stockfish_losses_with_black())
        table.cell(3,3).text = str(GC.stockfish_draws_with_black())


        doc.add_heading('2.2.2 Graphs', level=2)
        doc.add_paragraph('Here you can see a graph showing how many games were still going after x moves.')
        plt.clf()
        plt.figure(figsize=(6, 4))
        plt.plot(GC.stockfish_move_list_white(), color='red')
        plt.plot(GC.stockfish_move_list_black(), color='black')
        plt.plot(GC.stockfish_move_list_wins(), color='green')
        plt.xlabel('Move Number')
        plt.ylabel('Remaining Games')
        plt.legend(['White', 'Black', 'Wins'])
        plt.savefig('Plots/my_plot2.png')
        doc.add_picture('Plots/my_plot2.png')
        doc.add_paragraph(' ')

        doc.add_paragraph('Needed another graph to show the losses since there were so few.')
        plt.clf()
        plt.figure(figsize=(6, 4))
        plt.plot(GC.stockfish_move_list_losses(), color='red')
        plt.xlabel('Move Number')
        plt.ylabel('Remaining Games')
        plt.legend(['Losses'])
        plt.savefig('Plots/my_plot3.png')
        doc.add_picture('Plots/my_plot3.png')
        doc.add_paragraph(' ')

        #3.1 Mean and Standard Deviation
        doc.add_heading('3.1 Mean and Standard Deviation for Stockfish', level=1)
        doc.add_paragraph('Here you can see the mean and standard deviation of the games played')
        doc.add_heading('3.1.1 statistics', level=2)
        table = doc.add_table(rows=3, cols=6)
        table.style = 'Table Grid'
        for row in table.rows:
            row.height =docx.shared.Cm(1)
        table.cell(1,0).text = 'Mean'
        table.cell(2,0).text = 'Standard Deviation'
        table.cell(0,1).text = 'All games'
        table.cell(0,2).text = 'Played as white'
        table.cell(0,3).text = 'Played as black'
        table.cell(0,4).text = 'Games won'
        table.cell(0,5).text = 'Games lost'
        table.cell(1,1).text = str(GC.stockfish_mean())
        table.cell(1,2).text = str(GC.stockfish_mean_white())
        table.cell(1,3).text = str(GC.stockfish_mean_black())
        table.cell(1,4).text = str(GC.stockfish_mean_wins())
        table.cell(1,5).text = str(GC.stockfish_mean_losses())
        table.cell(2,1).text = str(GC.stockfish_SD())
        table.cell(2,2).text = str(GC.stockfish_SD_white())
        table.cell(2,3).text = str(GC.stockfish_SD_black())
        table.cell(2,4).text = str(GC.stockfish_SD_wins())
        table.cell(2,5).text = str(GC.stockfish_SD_losses())

        print('Word document created')
        doc.save('ChessGames.docx')
 
 
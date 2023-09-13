from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


class DocumentWriter:

    def __init__(self, title, filename):
        self.document = Document()
        self.filename = filename
        self.current_section = None

        # Change font size
        self.document.styles['Normal'].font.name = 'Times New Roman'
        # Increase font size for all headings
        self.document.styles['Heading 1'].font.size = Pt(20)

        self.document.add_heading(
            title, 0).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_heading(
            "TPK4186 - Advanced Tools for Performance Engineering \n\n Assingment 3: Wafer Production Line").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_heading(
            '\nCreated by: Christian G Kartveit & Skjalg Nysaeter', 2).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def getDocument(self) -> Document:
        return self.document

    def addHeading(self, text, level):
        if not self.current_section:
            self.addSection()
        self.getDocument().add_heading(text, level)
        self.save()

    def getSection(self):
        return self.current_section

    def addParagraph(self, text):
        if not self.current_section:
            self.addSection()
        self.getDocument().add_paragraph(text)
        self.save()

    def addPicture(self, picture, size):
        if not self.current_section:
            self.addSection()
        self.getDocument().add_picture(picture, width=Inches(size))
        self.save()

    def addTable(self, data, headers=None):  # data is a list of lists
        if not self.current_section:
            self.addSection()
        rows = len(data)
        cols = len(data[0])
        table = self.getDocument().add_table(
            rows=rows+1, cols=cols, style='Light Shading')
        if headers:
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table.cell(i+1, j).text = cell
        self.addParagraph("\n")
        self.save()

    def addSection(self):
        self.getDocument().add_section()
        self.current_section = self.document.sections[-1]
        self.save()

    def save(self):
        self.document.save(str(self.filename) + '.docx')